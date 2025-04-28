# from utils.dify_api import audit_text

# def handle_uploaded_file(uploaded_file):
#     # 获取文件路径
#     file_path = f"uploaded_files/{uploaded_file.name}"

#     # 保存文件到本地
#     with open(file_path, "wb") as f:
#         f.write(uploaded_file.getbuffer())

#     # 读取文件内容
#     file_text = ""
#     if uploaded_file.type == "application/pdf":
#         # 这里用一些PDF阅读库来提取PDF内容，例如 PyPDF2 或 pdfplumber
#         import pdfplumber
#         with pdfplumber.open(file_path) as pdf:
#             file_text = "\n".join([page.extract_text() for page in pdf.pages])
#     elif uploaded_file.type == "application/msword" or uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         # 处理Word文件
#         import docx
#         doc = docx.Document(file_path)
#         file_text = "\n".join([para.text for para in doc.paragraphs])
#     elif uploaded_file.type == "application/vnd.ms-excel" or uploaded_file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
#         # 处理Excel文件
#         import pandas as pd
#         df = pd.read_excel(file_path)
#         file_text = df.to_string()
#     elif uploaded_file.type == "text/markdown":
#         # 处理Markdown文件
#         file_text = uploaded_file.getvalue().decode("utf-8")

#     # 调用审校API
#     audit_result = audit_text(file_text)

#     return file_path, audit_result

# from utils.dify_api import audit_text
# import pdfplumber
# import docx
# import pandas as pd

# def handle_uploaded_file(uploaded_file):
#     # 获取文件路径
#     file_path = f"uploaded_files/{uploaded_file.name}"

#     # 保存文件到本地
#     with open(file_path, "wb") as f:
#         f.write(uploaded_file.getbuffer())

#     # 读取文件内容
#     file_text = ""
#     if uploaded_file.type == "application/pdf":
#         # 这里用一些PDF阅读库来提取PDF内容，例如 PyPDF2 或 pdfplumber
#         with pdfplumber.open(file_path) as pdf:
#             file_text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text() is not None])
#     elif uploaded_file.type == "application/msword" or uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         # 处理Word文件
#         doc = docx.Document(file_path)
#         file_text = "\n".join([para.text for para in doc.paragraphs])
#     elif uploaded_file.type == "application/vnd.ms-excel" or uploaded_file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
#         # 处理Excel文件
#         df = pd.read_excel(file_path)
#         file_text = df.to_string()
#     elif uploaded_file.type == "text/markdown":
#         # 处理Markdown文件
#         file_text = uploaded_file.getvalue().decode("utf-8")

#     # 调用审校API
#     audit_result = audit_text(file_text)

#     return file_path, audit_result  # 返回文件路径和审校结果

# import os
# from utils.dify_api import audit_text

# def handle_uploaded_file(uploaded_file):
#     # 获取文件保存路径
#     file_path = f"uploaded_files/{uploaded_file.name}"

#     # 检查目录是否存在，如果不存在则创建
#     if not os.path.exists("uploaded_files"):
#         os.makedirs("uploaded_files")

#     # 保存文件到本地
#     with open(file_path, "wb") as f:
#         f.write(uploaded_file.getbuffer())

#     # 读取文件内容
#     file_text = ""
#     if uploaded_file.type == "application/pdf":
#         # 这里用一些PDF阅读库来提取PDF内容，例如 PyPDF2 或 pdfplumber
#         import pdfplumber
#         with pdfplumber.open(file_path) as pdf:
#             file_text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
#     elif uploaded_file.type == "application/msword" or uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         # 处理Word文件
#         import docx
#         doc = docx.Document(file_path)
#         file_text = "\n".join([para.text for para in doc.paragraphs])
#     elif uploaded_file.type == "application/vnd.ms-excel" or uploaded_file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
#         # 处理Excel文件
#         import pandas as pd
#         df = pd.read_excel(file_path)
#         file_text = df.to_string()
#     elif uploaded_file.type == "text/markdown":
#         # 处理Markdown文件
#         file_text = uploaded_file.getvalue().decode("utf-8")

#     # 调试：打印 file_text 内容
#     print(f"Extracted text from the file: {file_text[:200]}...")  # 打印前200个字符以便查看

#     # 检查 file_text 是否为空
#     if not file_text.strip():
#         raise ValueError(f"无法提取有效的文本内容，文件类型或内容可能无法识别：{uploaded_file.type}")

#     # 调用审校API
#     audit_result = audit_text(file_text)

#     return file_path, audit_result


import os
import streamlit as st
from utils.dify_api import audit_text

def handle_uploaded_file(uploaded_file):
    # 获取文件保存路径
    file_path = f"uploaded_files/{uploaded_file.name}"

    # 检查目录是否存在，如果不存在则创建
    if not os.path.exists("uploaded_files"):
        os.makedirs("uploaded_files")

    # 保存文件到本地
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # 读取文件内容
    file_text = ""
    with st.spinner("正在提取文件内容..."):
        if uploaded_file.type == "application/pdf":
            # 这里用一些PDF阅读库来提取PDF内容，例如 PyPDF2 或 pdfplumber
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                file_text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        elif uploaded_file.type == "application/msword" or uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # 处理Word文件
            import docx
            doc = docx.Document(file_path)
            file_text = "\n".join([para.text for para in doc.paragraphs])
        elif uploaded_file.type == "application/vnd.ms-excel" or uploaded_file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            # 处理Excel文件
            import pandas as pd
            df = pd.read_excel(file_path)
            file_text = df.to_string()
        elif uploaded_file.type == "text/markdown":
            # 处理Markdown文件
            file_text = uploaded_file.getvalue().decode("utf-8")

    # 调试：打印 file_text 内容
    st.write(f"提取的文件内容预览：{file_text[:200]}...")  # 显示提取的文件前200个字符

    # 检查 file_text 是否为空
    if not file_text.strip():
        st.error(f"无法提取有效的文本内容，文件类型或内容可能无法识别：{uploaded_file.type}")
        return None, "无法提取有效的文本内容"

    # 调用审校API
    with st.spinner("正在进行审校，请稍等..."):
        audit_result = audit_text(file_text)
    
    if not audit_result:
        st.error("审校结果为空，请检查文件内容或API设置。")
        return file_path, "审校失败"

    return file_path, audit_result
